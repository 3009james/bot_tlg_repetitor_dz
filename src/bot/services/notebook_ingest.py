from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

from src.bot.services.routerai_client import RouterAIClient


@dataclass
class NotebookDigest:
    title: str
    content_hash: str
    compact_context: str
    tokens_estimate: int


def extract_text_from_ipynb(raw_bytes: bytes) -> tuple[str, str]:
    try:
        payload = json.loads(raw_bytes.decode("utf-8-sig"))
    except Exception as exc:
        raise ValueError("Invalid .ipynb JSON content") from exc
    cells = payload.get("cells", [])
    chunks: list[str] = []
    for cell in cells:
        source = cell.get("source", [])
        if isinstance(source, list):
            text = "".join(source)
        else:
            text = str(source)
        text = text.strip()
        if text:
            chunks.append(text)
    all_text = "\n\n".join(chunks)
    title = chunks[0].split("\n")[0][:100] if chunks else "Notebook lesson"
    return title, all_text


def extract_text_from_txt(raw_bytes: bytes) -> tuple[str, str]:
    for enc in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            text = raw_bytes.decode(enc)
            break
        except Exception:
            text = ""
    if not text.strip():
        raise ValueError("Invalid .txt content")
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    title = lines[0][:100] if lines else "Text lesson"
    return title, text


def extract_text_from_docx(raw_bytes: bytes) -> tuple[str, str]:
    # Legacy binary Word format (.doc) often gets mislabeled as .docx.
    if raw_bytes[:4] == b"\xD0\xCF\x11\xE0":
        raise ValueError("This looks like legacy .doc. Please save as .docx or PDF")

    try:
        doc = Document(BytesIO(raw_bytes))
        chunks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        # Include text from tables too, it is common in lesson handouts.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        chunks.append(cell_text)
        if chunks:
            text = "\n\n".join(chunks)
            title = chunks[0][:100]
            return title, text
    except Exception:
        pass

    # Fallback parser for some non-standard but valid docx packages.
    try:
        with ZipFile(BytesIO(raw_bytes)) as zf:
            xml_data = zf.read("word/document.xml")
        root = ET.fromstring(xml_data)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        nodes = root.findall(".//w:t", ns)
        chunks = [((node.text or "").strip()) for node in nodes if (node.text or "").strip()]
    except Exception as exc:
        raise ValueError("Invalid .docx content") from exc

    if not chunks:
        raise ValueError(".docx has no readable text")
    text = "\n\n".join(chunks)
    title = chunks[0][:100]
    return title, text


def extract_text_from_pdf(raw_bytes: bytes) -> tuple[str, str]:
    try:
        reader = PdfReader(BytesIO(raw_bytes))
    except Exception as exc:
        raise ValueError("Invalid .pdf content") from exc
    chunks: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if text:
            chunks.append(text)
    if not chunks:
        raise ValueError(".pdf has no readable text")
    all_text = "\n\n".join(chunks)
    title = chunks[0].split("\n")[0][:100]
    return title, all_text


def extract_text_from_material(filename: str, raw_bytes: bytes) -> tuple[str, str]:
    ext = Path(filename).suffix.lower()
    if ext == ".ipynb":
        return extract_text_from_ipynb(raw_bytes)
    if ext == ".txt":
        return extract_text_from_txt(raw_bytes)
    if ext == ".docx":
        return extract_text_from_docx(raw_bytes)
    if ext == ".pdf":
        return extract_text_from_pdf(raw_bytes)
    raise ValueError("Unsupported file format. Use .txt, .docx or .pdf")


async def build_material_digest(filename: str, raw_bytes: bytes, llm: RouterAIClient) -> NotebookDigest:
    title, raw_text = extract_text_from_material(filename, raw_bytes)
    content_hash = hashlib.sha256(raw_bytes).hexdigest()
    compact_context = await llm.summarize_lesson(raw_text)
    tokens_estimate = max(1, len(compact_context) // 4)
    return NotebookDigest(
        title=title,
        content_hash=content_hash,
        compact_context=compact_context[:20000],
        tokens_estimate=tokens_estimate,
    )


async def build_notebook_digest(raw_bytes: bytes, llm: RouterAIClient) -> NotebookDigest:
    # Backward compatibility for old bot flow.
    return await build_material_digest("lesson.ipynb", raw_bytes, llm)
